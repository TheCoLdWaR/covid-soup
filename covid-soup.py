from bs4 import BeautifulSoup
import requests
import pandas as pd
import xlwt
from xlwt import Workbook
import docx
from docx.shared import Pt
import csv

url = "https://www.worldometers.info/coronavirus/"

request = requests.get(url).text

soup = BeautifulSoup(request,"lxml")

# General words
total_cases = "Total Cases : "
new_cases  = "New Cases : "
total_deaths = "Total Deaths : "
new_deaths = "New Deaths : "
total_recovered = "Total Recovered : "
new_recovered = "New Recovered : "
active_cases = "Active Cases : "
serious_critical = "Serious, Critical : "
currently_infected_patients = "Current Infected Patients : "
in_mild_condition = "in Mild Condition : "
serious_or_critical = "Serious or Critical : "
population = "Population : "

words = [total_cases,new_cases,total_deaths,new_deaths,total_recovered,new_recovered,active_cases,serious_critical]

div = soup.find_all("div", attrs = {"class":"maincounter-number"})
# Total Statistics!
total_corona_virus_number = "Total Coronavirus Cases : " + div[0].text.lstrip()
total_corona_virus_deaths = "Total Coronavirus Deaths : " + div[1].text.lstrip()
total_corona_virus_recovered = "Total Coronavirus Recovered : " + div[2].text.lstrip()

tbody = soup.find("tbody")
europe = tbody.find("tr", attrs = {"data-continent":"Europe"})
nortAmerica = tbody.find("tr", attrs = {"data-continent":"North America"})
asia = tbody.find("tr", attrs = {"data-continent":"Asia"})
southAmerica = tbody.find("tr", attrs = {"data-continent":"South America"})
africa = tbody.find("tr", attrs = {"data-continent":"Africa"})
oceania = tbody.find("tr", attrs= {"data-continent":"Australia/Oceania"})

# main options
def options():

    whichOption = input("Choose One:\n------\n1- Continents\n2- Countries\n3- Which Countries\n4- World Total Stats\n")
    if whichOption == "1":
        writeContinent()
    elif whichOption == "2":
        writeCountry()
    elif whichOption == "3":
        listCountries()
    elif whichOption == "4":
        worldTotalStats()
    else:
        print("Wrong Choise or missing!!")

def listCountries():
    # import countries

    url_countries = "https://www.worldometers.info/geography/alphabetical-list-of-countries/"

    r_countries = requests.get(url_countries).text

    countries = []

    soup = BeautifulSoup(r_countries,"lxml")

    tbody = soup.find("tbody")

    tr_tags = tbody.find_all("tr")

    country_names = soup.find_all("td", attrs = {"style":"font-weight: bold; font-size:15px"})

    file = open("country_list.txt","w+")

    for datas in country_names:
        countries.append(datas.text)

    #  - Some fixes -
    countries[26] = "Burkina-Faso"
    countries[21] = "Bosnia-and-herzegovina"
    countries[5] = "Antigua-and-barbuda"
    countries[133] = "State-of-palestine"
    countries[185] = "UK"
    countries[186] = "USA"
    countries[44] = "Czechia"
    countries[157] = "Swaziland"
    countries[56] = "Swaziland"
    countries[33] = "Central-African-Republic"
    countries[39] = "Congo"
    countries[45] = "Democratic-Republic-of-the-Congo"
    countries[49] = "Dominican-Republic"
    countries[52] = "El-Salvador"
    countries[53] = "Equatorial-Guinea"
    countries[73] = "Holy-See"
    countries[107] = "Marshall-islands"
    countries[118] = "Myanmar"
    countries[123] = "New-Zealand"
    countries[127] = "North Korea"
    countries[128] = "Macedonia"
    countries[135] = "Papua New Guinea"
    countries[145] = "Saint-kitts-and-nevis"
    countries[146] = "saint-lucia"
    countries[149] = "san-marino"
    countries[150] = "sao-tome-and-principe"
    countries[151] = "saudi-arabia"
    countries[155] = "sierra-leone"
    countries[159] = "solomon-islands"
    countries[161] = "south-africa"
    countries[163] = "south-sudan"
    countries[162] = "South-Korea"
    countries[165] = "Sri-Lanka"
    countries[177] = "trinidad-and-tobago"
    countries[184] = "united-arab-emirates"

    print()

    for x in range(0,195,1):

        file.write(countries[x] + "\n")

    file.close()
    print("country_list.txt printing completed succesfully!")

def writeContinent():

    print("Continents : \n1- Europe\n2- North America\n3- Asia\n4- South America\n5- Africa\n6- Oceania\n")

    continent_name_input = input("Enter Continent number, between 1-6\n")

    if continent_name_input == "1":

        #Europe
        europe_td_tags = europe.find_all("td")

        europe_statistics = []
        for datas in europe_td_tags[2:-12]:
            europe_statistics.append(datas.text)
            # print(datas.text)

        europe_dataframe = pd.DataFrame(europe_statistics, words ,columns=["Last 24 Hours"])

        print(europe_dataframe)
        writecontinentOptions("Europe", europe_statistics)

    elif continent_name_input == "2":

        # North America
        nortAmerica_td_tags = nortAmerica.find_all("td")

        nortAmerica_statistics = []
        for datas in nortAmerica_td_tags[2:-12]:
            nortAmerica_statistics.append(datas.text)

        northAmerica_dataframe = pd.DataFrame(nortAmerica_statistics, words,columns=["North America Last 24 Hours"])

        print(northAmerica_dataframe)
        writecontinentOptions("North America", nortAmerica_statistics)

    elif continent_name_input == "3":

        # Asia
        asia_td_tags = asia.find_all("td")

        asia_statistics = []
        for datas in asia_td_tags[2:-12]:
            asia_statistics.append(datas.text)

        asia_dataframe = pd.DataFrame(asia_statistics, words, columns=["Asia Last 24 Hours"])

        print(asia_dataframe)

        writecontinentOptions("Asia", asia_statistics)

    elif continent_name_input == "4":

        # South America
        southAmerica_td_tags = southAmerica.find_all("td")

        southAmerica_statistics = []
        for datas in southAmerica_td_tags[2:-12]:
            southAmerica_statistics.append(datas.text)
            print(datas.text)

        southAmerica_dataframe = pd.DataFrame(southAmerica_statistics, words, columns=["South America Last 24 Hours"])

        print(southAmerica_dataframe)
        writecontinentOptions("South America", southAmerica_statistics)

    elif continent_name_input == "5":

        # Africa
        africa_td_tags = africa.find_all("td")

        africa_statistics = []
        for datas in africa_td_tags[2:-12]:
            africa_statistics.append(datas.text)

        africa_dataframe = pd.DataFrame(africa_statistics, words, columns=["Africa Last 24 Hours"])

        print(africa_dataframe)
        writecontinentOptions("Africa", africa_statistics)

    elif continent_name_input == "6":

        # Oceania
        oceania_td_tags = oceania.find_all("td")

        oceania_statistics = []
        for datas in oceania_td_tags[2:-12]:
            oceania_statistics.append(datas.text)

        oceania_dataframe = pd.DataFrame(oceania_statistics, words, columns=["Oceania/Australia Last 24 Hours"])

        print(oceania_dataframe)

        writecontinentOptions("Oceania", oceania_statistics)

    else:
        print("Wrong choise or missing !!")


#worldTotalStats
def worldTotalStats():
    url = "https://www.worldometers.info/coronavirus/"

    r = requests.get(url).text

    soup = BeautifulSoup(r,"lxml")

    world_total_cases_data = soup.find("span", attrs = {"style":"color:#aaa"}).text
    world_total_deaths_data = soup.find_all("div", attrs = {"style":"margin-top:15px"})
    world_total_recovered_data = soup.find("div", attrs = {"style":"color:#8ACA2B "}).text

    world_total_cases = "World Total Cases : " + world_total_cases_data
    world_total_deaths = "World Total Deaths :" + " " + world_total_deaths_data[1].text.replace('Deaths:','').strip()
    world_total_recovered = "World Total Recovered Data : " + " " + world_total_recovered_data.strip()

    # print(world_total_cases)
    print(world_total_cases)
    print(world_total_deaths)
    print(world_total_recovered)

    mydoc = docx.Document()

    style = mydoc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(15)

    mydoc.add_heading("World Total Cases", 0)
    mydoc.add_paragraph(world_total_cases)
    mydoc.add_paragraph(world_total_deaths)
    mydoc.add_paragraph(world_total_recovered)

    file_name = input("Set your file name ...(Example : covid_news.docx)\n")

    mydoc.save(file_name)

def writeOptions(country_name_input,country_total_cases,country_total_deaths,country_total_recovered,country_new_datas
                ,maincounter_datas,news_li):

                def writetoDocx(country_name_input,country_total_cases,
                            country_total_deaths,country_total_recovered,country_new_datas):

                    mydoc = docx.Document()
                    style = mydoc.styles['Normal']
                    font = style.font
                    font.name = 'Arial'
                    font.size = Pt(15)

                    mydoc.add_heading(country_name_input + " Covid-19 Statistics",0)
                    mydoc.add_paragraph(country_total_cases)
                    mydoc.add_paragraph(country_total_deaths)
                    mydoc.add_paragraph(country_total_recovered)
                    mydoc.add_paragraph(country_new_datas)

                    file_name = input("Set your file name ...\n(Example : covid_news.docx , covid_news.xls)\n")

                    mydoc.save(file_name)

                    print("File Name : " +  file_name + " Printing completed succesfully !")

                def writeToExcel(country_name_input,maincounter_datas,news_li,country_total_cases,country_total_deaths,country_total_recovered):

                    wb = Workbook()

                    # cds = covid data sheet
                    cds = wb.add_sheet("covid")

                    font = xlwt.Font()
                    font = height = 15

                    cds.write(0,0,country_name_input + " Covid-19 Statistics")

                    cds.write(1,0,country_name_input + " "  + total_cases)
                    cds.write(1,1,maincounter_datas[0].text.lstrip())

                    cds.write(2,0,country_name_input + " "  + total_deaths)
                    cds.write(2,1,maincounter_datas[1].text.lstrip())

                    cds.write(3,0,country_name_input+" " + total_recovered)
                    cds.write(3,1,maincounter_datas[2].text.lstrip())

                    cds.write(4,0,"Last 24 Hours :" + " ")
                    cds.write(4,1,news_li[0].text.replace('[source]',''))

                    file_name = input("Set your file name ...\n(Example : covid_news.docx , covid_news.xls)\n")

                    wb.save(file_name)

                    print("File Name : "+file_name + " Printing completed succesfully !")

                def writetoCSV(country_name_input,country_total_cases,country_total_deaths,country_total_recovered,country_new_datas):

                    file_name = input("Set your file name ...(Example : covid_news.docx)\n")

                    with open(file_name,mode="w") as employee_file:
                        employee_file = csv.writer(employee_file, delimiter=',',quoting=csv.QUOTE_MINIMAL)

                        employee_file.writerow([country_name_input,country_total_cases,country_total_deaths,country_total_recovered,country_new_datas])

                        print("Printing completed succesfully!")

                writeOptionsInput = input("\n1-Write to Excel file\n2-Write to .docx file\n3-Write to csv file\n")

                #writetoExcel
                if writeOptionsInput == "1":

                    writeToExcel(country_name_input,maincounter_datas,news_li,country_total_cases,country_total_deaths,country_total_recovered)

                #writeTodocx
                elif writeOptionsInput == "2":

                    writetoDocx(country_name_input,country_total_cases,country_total_deaths,country_total_recovered,
                                    country_new_datas)
                elif writeOptionsInput == "3":

                    writetoCSV(country_name_input,country_total_cases,country_total_deaths,country_total_recovered,
                                    country_new_datas)

def writeCountry():

    country_name_input = input("Enter a country name :\n")

    try:

        # Search specific country data.
        url2 = "https://www.worldometers.info/coronavirus/country/" + country_name_input
        request2 = requests.get(url2).text
        soup = BeautifulSoup(request2,"lxml")

        maincounter_datas = soup.find_all("div", attrs = {"class":"maincounter-number"})

        news_li = soup.find_all("li", attrs = {"class":"news_li"})

        active_cases_data1 = soup.find("div", attrs = {"class":"number-table-main"})
        active_cases_data2 = soup.find("span", attrs = {"class":"number-table"})
        active_cases_data3 = soup.find("div", attrs = {"style":"float:right; text-align:center"})

        country_total_cases = country_name_input + " "  + total_cases + maincounter_datas[0].text.strip()
        country_total_deaths = country_name_input + " "  + total_deaths + maincounter_datas[1].text.strip()
        country_total_recovered = country_name_input + " " + total_recovered + maincounter_datas[2].text.strip()
        country_new_datas = "Last 24 Hours :" + " " + news_li[0].text.replace('[source]','')

        print(country_total_cases)
        print(country_total_deaths)
        print(country_total_recovered)
        print(country_new_datas) #buraya kadar sıkınt yoq

        writeOptions(country_name_input,country_total_cases,country_total_deaths,
                country_total_recovered,country_new_datas,maincounter_datas,news_li)


    except Exception as error:
        print(error)
        print("No up-to-date data for or wrong country name " + country_name_input)

def writecontinentOptions(continent_name_input,statistic=[]):


    print("--------------------")
    print("--------------------")
    input_option = input("Choose one\n1-Write to .docx\n2-Write to .xls\n")

    if input_option == "1":
        #writecontinentdataDocx

            mydoc_continent = docx.Document()
            style = mydoc_continent.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(15)

            mydoc_continent.add_heading(continent_name_input + " Covid-19 Statistics",0)
            mydoc_continent.add_paragraph(total_cases + statistic[0])
            mydoc_continent.add_paragraph(new_cases + statistic[1])
            mydoc_continent.add_paragraph(total_deaths + statistic[2])
            mydoc_continent.add_paragraph(new_deaths + statistic[3])
            mydoc_continent.add_paragraph(total_recovered + statistic[4])
            mydoc_continent.add_paragraph(new_recovered + statistic[5])
            mydoc_continent.add_paragraph(active_cases + statistic[6])
            mydoc_continent.add_paragraph(serious_critical + statistic[7])


            file_name = input("Set your file name ...(Example : covid_news.docx)\n")
            mydoc_continent.save(file_name)
            print(file_name + "Printing completed succesfully!")

    elif input_option == "2":

        wb1 = Workbook()

        # cds = covid data sheet
        cds = wb1.add_sheet("continent_name_input")

        font = xlwt.Font()
        font = height = 15

        cds.write(0,0,continent_name_input + " Covid-19 Statistics")

        cds.write(1,0,continent_name_input +  " " + total_cases)
        cds.write(1,1,statistic[0])

        cds.write(2,0,continent_name_input + " " + new_cases)
        cds.write(2,1,statistic[1])

        cds.write(3,0,continent_name_input + " " + total_deaths)
        cds.write(3,1,statistic[2])

        cds.write(4,0,continent_name_input + " " + new_deaths)
        cds.write(4,1,statistic[3])

        cds.write(5,0,continent_name_input + " " + total_recovered)
        cds.write(5,1,statistic[4])

        cds.write(6,0,continent_name_input + " " + new_recovered)
        cds.write(6,1,statistic[5])

        cds.write(7,0,continent_name_input + " " + active_cases)
        cds.write(7,1,statistic[6])

        cds.write(8,0,continent_name_input + " " + serious_critical)
        cds.write(8,1,statistic[7])

        file_name = input("Set your file name ...(Example : covid_news.docx)\n")

        wb1.save(file_name)

        print(file_name + "Printing completed succesfully!")

    elif input_option == "3":
        # print("Wrong Choice or missing!")
        print("merhaba")

    else :
        pass



options()
